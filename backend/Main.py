
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Dict, List, Optional, Any
import os
import shutil
import uuid
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import base64
import json
import requests
from googletrans import Translator
import tempfile
import pypandoc
from dotenv import load_dotenv
import time
from pathlib import Path
import cv2
import numpy as np
from functools import lru_cache
import re
from docx import Document
from docx.shared import Pt, Inches
from fastapi import HTTPException
from fastapi.responses import FileResponse
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT



# Load environment variables
load_dotenv()

# Set Tesseract executable path - update this path to match your Tesseract installation
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

app = FastAPI(title="Fox Mandal OCR-AI API")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify the exact frontend origin
    allow_credentials=True,
    allow_methods=["*"],
    
)

# Create necessary directories
os.makedirs("uploads", exist_ok=True)
os.makedirs("images", exist_ok=True)
os.makedirs("temp", exist_ok=True)
os.makedirs("outputs", exist_ok=True)

# Environment variables
API_KEY = "rGcsjveDY8gXRWIJfuTJUBGNKE17BPvI6aTKWXVDtcF3"
PROJECT_ID = "23d2c61b-83d2-4985-8bc4-afe65f32bef8"


SECTION_PROMPTS = {
    "I. DESCRIPTION OF THE LANDS": """
    You are a legal expert analyzing land documents. Extract and format the land description exactly as:

    Example Output:
    Survey No. 46/1, measuring to an extent of 10 acres 11 guntas, situated at Harlapura village, Betageri hobli, Gadag taluk, and Gadag district.

    Instructions:
    1. Extract only survey number, extent, and location details
    2. Format exactly as shown in the example
    3. Include all location details (village, hobli, taluk, district)
    4. Do not include any additional text or explanations
    """,

    "II. LIST OF THE DOCUMENT REVIEWED": """
    You are a legal expert. Extract and list all documents reviewed exactly as (maximum 25 document):

    Example Output:
    Serial No.
    Description of Documents

    1.
    Record of Tenancy and Crops for the period 1987-88 to 2003-04, issued by the office of Tahsildar, Gadag taluk

    2.
    Mutation Register Extract bearing No. 120/2003-04, issued by the office of Tahsildar, Gadag taluk

    3.
    Record of Tenancy and Crops for the period 2004-05 to 2015-16, issued by the office of Tahsildar, Gadag taluk

    Instructions:
    1. Use sequential numbering (1., 2., 3., etc.)
    2. Each document description must include all relevant details
    3. Format exactly as shown in the example
    4. Maintain proper spacing between entries
    """,

    "III. DEVOLUTION OF TITLE": """
    You are a legal expert analyzing title chain. Format the output exactly as:

    Example Output:
    SL No. | SURVEY No. | EXTENT | EXTENT OF KHARAB LAND | OWNER/S
    ------|------------|---------|----------------------|---------
    1. | 46/1 | 10 | 11 | 00 | 00 | Mr. Chandrashekar s/o. Shivaji Halalli

    Upon perusal of the documents furnished to us,

    1. It is learnt from the Record of Tenancy and Crops for the period 1987-88 to 2003-04, issued by the office of Tahsildar, Gadag taluk, that the name of Mr. Shivaji s/o. Ramappa Halalli is recorded as owner in possession of land bearing Survey No. 46, measuring an extent of 20 acres 22 guntas.

    2. It is observed from the Mutation Register Extract bearing No. 120/2003-04, issued by the office of Tahsildar, Gadag taluk, that Mr. Shivaji s/o. Ramappa Halalli has availed loan amount of Rs. 50,000 and mortgaged the land bearing Survey No. 46 measuring 20 acres 22 guntas in favour of Vyavasaya Seva Sahakari Bank, Harlapura.

    Instructions:
    1. Start with the table showing current ownership
    2. Follow with numbered list of ownership history
    3. Include all relevant dates and document references
    4. Format exactly as shown in the example
    """,

    "IV. ENCUMBRANCE CERTIFICATE": """
    You are a legal expert reviewing encumbrances. Format the output exactly as:

    Example Output:
    1. Encumbrance Certificate for the period from 01.04.1985 to 31.03.2004, issued by the office of Sub-Registrar, Gadag, with regard to Survey No. 46 measuring to an extent of 20 acres 22 guntas, does not reflect any transactions.

    2. Encumbrance Certificate for the period from 01.04.2004 to 12.08.2024, issued by the office of Sub-Registrar, Gadag, with regard to Survey No. 46 measuring to an extent of 20 acres 22 guntas, reflects the entries as follows:
    Sl. No. | Transactions | Document No | Dated | Remark
    --------|-------------|-------------|-------|--------
    1. | Partition Deed | document No. GDG-1-09812/2017-18 | 19.01.2018 | Nil
    2. | Mortgage Deed | document No. GDG-1 Part-V-00110/2017-18 | 07.12.2017 | Nil
    3. | Gift Deed | document No. GDG-1-10167/2014-15 | 10.03.2015 | Nil

    Instructions:
    1. List each encumbrance certificate separately
    2. Include all relevant details and dates
    3. Format transactions in table format when present
    4. Format exactly as shown in the example
    """,

    "V. OTHER OBSERVATIONS": """
    You are a legal expert making observations. Format the output exactly as:

    Example Output:
    (i) ALL THAT PIECE AND PARCEL of the Agricultural land bearing Survey No. 46/1 measuring 10 acres 11 guntas, situated at Harlapura village, Gadag taluk, Betageri hobli, Gadag district and bound on:
    East by : Survey No. 47
    West by : Survey No. 43
    North by : Survey No. 45
    South by : Survey No. 46/2.
    [Boundaries are ascertained from the Tippani, PT sheet/Ghat plot]

    (ii) RESTRICTIONS ON TRANSFERABILITY
    a. Land Ceiling: - The Measurement of Schedule Property falls within the prescribed limit provided under Section 63 of Karnataka Land Reforms Act.
    b. Minor's interest: - NO
    c. Grant/Inam Lands: - NO

    (iii) ENDORSMENTS:
    Note: PTCL, Nil Tenancy and Nil Acquisition endorsement issued by the concerned authority.

    (iv) FAMILY TREE OF THE CURRENT LANDOWNERS
    1. It is learnt from the Notarized Genealogical Tree dated 10.09.2024, 02.01.2025, declared by Mr. Chandrashekar s/o. Shivaji Halalli.

    Husband: - Mr. Chandrashekar alias Chandrashekarappa s/o. Shivaji Halalli (50 years)
    Wife: - Mrs. Neelamma (45 years)

    1. Mrs. Kaveri w/o. Manjappa Honalli (27 years)
    2. Mrs. Bheemavva w/o. Gavisiddappa Arera (24 years)
    3. Ms. Lakshmavva d/o. Chandrashekar Halalli (23 years) unmarried
    4. Ms. Yallamma d/o. Chandrashekar Halalli (19 years) unmarried
    5. Master. Venkappa alias Yankappa s/o. Chandrashekar Halalli (15 years)

    Instructions:
    1. Include all sections (i) through (iv)
    2. Format exactly as shown in the example
    3. Maintain proper spacing and indentation
    """,

    "VI. INDEPENDENT VERIFICATIONS": """
    You are a legal expert verifying documents. Format the output exactly as:

    Example Output:
    (i) Sub-Registrar Search's: The Sub-Registrar Search Report, issued by Mr. B.P.Gubber Advocate and the same is attached to this report as an annexure.

    (ii) Revenue Records Search: The Revenue Search Report, issued by Mr. B.P.Gubber Advocate and the same is attached to this report as an annexure.

    Instructions:
    1. List each verification separately
    2. Include all relevant details and findings
    3. Format exactly as shown in the example
    """,

    "VII. LITIGATION SEARCH RESULTS": """
    You are a legal expert analyzing litigation. Format the output exactly as:

    Example Output:
    (i) The Litigation Search Report, issued by Mr. B.P.Gubber Advocate and the same is attached to this report as an annexure.

    (ii) The PACL Land scam Search Report, issued by Mr. B.P.Gubber Advocate and the same is attached to this report as an annexure.

    Instructions:
    1. List each search result separately
    2. Include all relevant details and findings
    3. Format exactly as shown in the example
    """,

    "VIII. SPECIAL CATEGORY LANDS": """
    You are a legal expert categorizing land. Format the output exactly as:

    Example Output:
    Upon perusal of documents scrutinized above, it is found that the schedule property DOES NOT come under the purview of SC/ST/Minors/Inam/Grant lands or any land under Special Categories.

    Instructions:
    1. Provide clear categorization
    2. Include all relevant details
    3. Format exactly as shown in the example
    """,

    "IX. OPINION AND RECOMMENDATION": """
    You are a legal expert providing opinion. Format the output exactly as:

    Example Output:
    Upon review and scrutiny of the documents furnished to us and based on independent searches by Mr.B.P. Gubber Advocate, we are of the opinion that, Mr. Chandrashekar s/o. Shivaji Halalli is the absolute owner having valid, clear and marketable title, with respect to land bearing Survey No. 46/1 measuring to an extent of 10 acres 22 guntas, situated at Harlapura village, Gadag taluk, betageri hobli, Gadag district.

    Following person are to be joined as signatories in the future Deed/s:
    Sl. No. | Owner/s or Khatedars or Co-owners | Sl.No | Family Members
    --------|-----------------------------------|-------|---------------
    1 | Mr. Chandrashekar s/o. Shivaji Halalli | 1 | Mrs. Neelamma w/o. Chandrashekar Halalli
    | | | 2 | Mrs. Kaveri w/o. Manjappa Honalli
    | | | 3 | Mrs. Bheemavva w/o. Gavisiddappa Arera
    | | | 4 | Ms. Lakshmavva d/o. Chandrashekar Halalli
    | | | 5 | Ms. Yallamma d/o. Chandrashekar Halalli
    | | | 6 | Master. Venkappa alias Yankappa (15 years) M/g father Mr. Chandrashekar Halalli

    However, the same is subject to the following documents / clarification:
    1. Discharge of mortgage created vide MR No. 120/2003-04 in favour of Vyavasaya Seva Sahakari Bank, Harlapura.
    2. Discharge of mortgage created vide mortgage deed registered on 06.01.2018 as document No. GDG-1 Part-V-00110/2017-18.
    3. Latest Property Tax paid receipt.
    4. Family Tree of Mr. Chandrashekar s/o. Shivaji Halalli issued by the office of Tahsildar.
    5. PTCL, Nil Tenancy and Nil Acquisition endorsement issued by the concerned authority.
    6. 11 e sketch shall be provided before execution of sale deed and alienation sketch shall be provided before NA application for lease area.
    7. The subject land is an Agricultural Land hence before using it for Non-Agricultural purpose, Company shall obtain Conversion Order under Sec 95 of KLR Act 1964.

    Instructions:
    1. Include opinion statement
    2. List all required signatories in table format
    3. List all requirements and clarifications
    4. Format exactly as shown in the example
    """,

    "X. CONTACT DETAILS": """
    You are a legal expert providing contact information. Format the output exactly as:

    Example Output:
    If any clarification in relation to this Report is required, please contact:

    Prashantha Kumar S. T
    Senior Partner
    Fox Mandal & Associates
    "FM House"
    6/12, Primrose Road
    Bangalore 560 025
    Phone  : +91 80 2559 5911
    Mobile : +91 98801 62142
    e-mail : prashantha.kumar@foxmandal.in

    Instructions:
    1. Include all contact details
    2. Format exactly as shown in the example
    3. Maintain proper spacing and alignment
    """
}

class ProcessingStatus(BaseModel):
    session_id: str
    status: str
    message: str
    progress: float
    current_stage: str
    total_pages: int
    processed_pages: int
    final_output: Optional[str] = None

class ProcessingResponse(BaseModel):
    session_id: str
    message: str

class PageData(BaseModel):
    page_number: int
    raw_text: str
    translated_text: str

class PageUpdateRequest(BaseModel):
    page_number: int
    edited_text: str
    
class ReportRequest(BaseModel):
    session_id: str
    client_name: Optional[str] = None

class DocumentSuggestion(BaseModel):
    name: str
    required: bool = True
    uploaded: bool = False

class DocumentRequest(BaseModel):
    chunk_text: str

# In-memory storage for process tracking
processing_status = {}

def get_ibm_access_token(api_key):
    """Get IBM WatsonX access token"""
    url = "https://iam.cloud.ibm.com/identity/token"
    headers = {"Content-Type": "application/x-www-form-urlencoded"}
    data = {
        "grant_type": "urn:ibm:params:oauth:grant-type:apikey",
        "apikey": api_key
    }
    response = requests.post(url, headers=headers, data=data)
    return response.json()["access_token"]

def preprocess_image(pil_image):
    """Preprocess image to improve OCR quality"""
    img = np.array(pil_image.convert("RGB"))
    img = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    img = cv2.resize(img, None, fx=1.5, fy=1.5, interpolation=cv2.INTER_LINEAR)
    img = cv2.fastNlMeansDenoising(img, h=30)
    kernel = np.array([[0, -1, 0],
                       [-1, 5,-1],
                       [0, -1, 0]])
    img = cv2.filter2D(img, -1, kernel)
    img = cv2.adaptiveThreshold(img, 255,
                                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                cv2.THRESH_BINARY, 35, 15)
    return Image.fromarray(img)

def extract_text_from_image(image: Image.Image):
    """Extract text from image using OCR"""
    try:
        # Preprocess image
        processed_img = preprocess_image(image)
        
        # Perform OCR with Tesseract
        extracted_text = pytesseract.image_to_string(processed_img, lang='kan+eng')
        return extracted_text
    except Exception as e:
        return f"[OCR failed: {str(e)}]"

def translate_text(text: str, src='kn', dest='en'):
    """Translate text from one language to another"""
    translator = Translator()
    try:
        translated = translator.translate(text, src=src, dest=dest).text
        return translated
    except Exception as e:
        return f"[Translation failed: {str(e)}]"

def chunk_text(text_dict: Dict[str, str], chunk_size=15):
    """Split text into manageable chunks for AI processing"""
    pages = list(text_dict.items())
    return [dict(pages[i:i + chunk_size]) for i in range(0, len(pages), chunk_size)]

def send_chunk_to_watsonx(chunk_text: str, access_token: str, section_prompt: str):
    """Send text chunk to WatsonX AI for processing"""
    url = "https://us-south.ml.cloud.ibm.com/ml/v1/text/generation?version=2024-01-15"
    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Authorization": f"Bearer {access_token}"
    }
    payload = {
        "input": section_prompt + "\n\nSourceDocument text:\n" + chunk_text,
        "parameters": {
            "decoding_method": "greedy",
            "max_new_tokens": 100,
            "min_new_tokens": 0,
            "stop_sequences": [],
            "repetition_penalty": 1
        },
        "model_id": "meta-llama/llama-3-3-70b-instruct",
        "project_id": PROJECT_ID
    }
    response = requests.post(url, headers=headers, json=payload)
    try:
        response.raise_for_status()
        result = response.json()
        if "results" in result:
            return result["results"][0]["generated_text"]
        else:
            return f"[WatsonX response error: {result.get('errors', 'Unknown error')}]"
    except requests.exceptions.HTTPError as http_err:
        return f"[WatsonX HTTP error: {str(http_err)} - Raw: {response.text}]"
    except Exception as e:
        return f"[WatsonX response error: {str(e)} - Raw: {response.text}]"

def analyze_image_quality(image: Image.Image) -> bool:
    """Analyze image quality using noise detection"""
    # Convert PIL image to numpy array
    img_array = np.array(image.convert("L"))
    
    # Apply Canny edge detection to find edges
    edges = cv2.Canny(img_array, 100, 200)
    
    # Calculate noise level using standard deviation of edges
    noise = np.std(edges)
    
    # Return True if quality is poor (high noise)
    return noise >= 70

def process_pdf(session_id: str, file_path: str, background_tasks: BackgroundTasks):
    """Process PDF file in background"""
    try:
        # Initialize status tracking
        processing_status[session_id] = {
            "status": "processing",
            "message": "Starting PDF processing",
            "progress": 0.0,
            "current_stage": "initialization",
            "total_pages": 0,
            "processed_pages": 0,
            "extracted_pages": {},
            "translated_pages": {},
            "edited_pages": {},
            "pdf_images": {},
            "poor_quality_pages": [],
            "final_output": None
        }
        
        # Create session directory for this processing job
        session_dir = os.path.join("temp", session_id)
        os.makedirs(session_dir, exist_ok=True)
        images_dir = os.path.join("images", session_id)
        os.makedirs(images_dir, exist_ok=True)
        
        # Update status
        processing_status[session_id].update({
            "message": "Opening PDF document",
            "progress": 0.05,
            "current_stage": "pdf_loading"
        })
        
        # Extract pages and preload images
        extracted_pages = {}
        translated_pages = {}
        pdf_images = {}
        poor_quality_pages = []
        
        with fitz.open(file_path) as doc:
            total_pages = len(doc)
            processing_status[session_id]["total_pages"] = total_pages
            
            for page_num in range(total_pages):
                # Update progress
                processing_status[session_id].update({
                    "message": f"Processing page {page_num+1} of {total_pages}",
                    "progress": 0.1 + (0.7 * (page_num / total_pages)),
                    "current_stage": "ocr_translation",
                    "processed_pages": page_num
                })
                
                # Get page
                page = doc.load_page(page_num)
                
                # Render page to image for OCR
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_bytes = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_bytes))
                
                # Check image quality
                if analyze_image_quality(img):
                    poor_quality_pages.append(page_num + 1)
                
                # Save image for future reference
                image_path = os.path.join(images_dir, f"page_{page_num+1}.png")
                img.save(image_path)
                
                # Convert to base64 for frontend
                img_base64 = base64.b64encode(img_bytes).decode()
                pdf_images[page_num] = img_base64
                
                # Perform OCR
                extracted_text = extract_text_from_image(img)
                extracted_pages[f"Page {page_num+1}"] = extracted_text
                
                # Translate text
                translated_text = translate_text(extracted_text, src='kn', dest='en')
                translated_pages[f"Page {page_num+1}"] = translated_text
                
                # Add small delay to avoid overwhelming resources
                time.sleep(0.1)
        
        # Update status
        processing_status[session_id].update({
            "message": "OCR and translation completed",
            "progress": 0.8,
            "current_stage": "completed",
            "processed_pages": total_pages,
            "extracted_pages": extracted_pages,
            "translated_pages": translated_pages,
            "edited_pages": {k: v for k, v in translated_pages.items()},
            "pdf_images": pdf_images,
            "poor_quality_pages": poor_quality_pages
        })
        
        # Save results to files for persistence
        with open(os.path.join(session_dir, "extracted_pages.json"), "w", encoding="utf-8") as f:
            json.dump(extracted_pages, f, ensure_ascii=False, indent=2)
            
        with open(os.path.join(session_dir, "translated_pages.json"), "w", encoding="utf-8") as f:
            json.dump(translated_pages, f, ensure_ascii=False, indent=2)
            
        with open(os.path.join(session_dir, "pdf_images.json"), "w", encoding="utf-8") as f:
            json.dump(pdf_images, f, ensure_ascii=False, indent=2)
            
        with open(os.path.join(session_dir, "poor_quality_pages.json"), "w", encoding="utf-8") as f:
            json.dump(poor_quality_pages, f, ensure_ascii=False, indent=2)
        
        # Final update
        processing_status[session_id].update({
            "status": "ready_for_review",
            "message": "PDF processing complete! Ready for quality review.",
            "progress": 1.0,
            "current_stage": "waiting_for_review"
        })
        
    except Exception as e:
        # Update status on error
        processing_status[session_id].update({
            "status": "error",
            "message": f"Error processing PDF: {str(e)}",
            "progress": 0,
            "current_stage": "error"
        })


        
def generate_report(session_id: str, client_name: Optional[str] = None):
    """Generate final report using WatsonX AI"""
    try:
        # Update status
        processing_status[session_id].update({
            "status": "generating_report",
            "message": "Starting report generation",
            "progress": 0.0,
            "current_stage": "starting_report"
        })

        session_dir = os.path.join("temp", session_id)
        output_dir = os.path.join("outputs", session_id)
        os.makedirs(output_dir, exist_ok=True)

        # Get edited pages
        edited_pages = processing_status[session_id]["edited_pages"]
        combined_text = "\n".join(edited_pages.values())

        # Get IBM WatsonX token
        token = get_ibm_access_token(API_KEY)

        # Initialize final_output string
        final_output = ""


        # Process each section separately
        section_results = {}
        total_sections = len(SECTION_PROMPTS)
        
        for i, (section_name, prompt) in enumerate(SECTION_PROMPTS.items()):
            # Update progress
            progress = 0.2 + (0.6 * (i / total_sections))
            processing_status[session_id].update({
                "message": f"Processing section {i+1} of {total_sections}: {section_name}",
                "progress": progress,
                "current_stage": f"processing_section_{i+1}"
            })

            # Send section-specific prompt to WatsonX
            result = send_chunk_to_watsonx(combined_text, token, prompt)
            section_results[section_name] = result

       
        for section_name, content in section_results.items():
            final_output += f"## {section_name}\n\n{content}\n\n"
         
        if client_name:
            final_output = final_output.replace("[Client Name]", client_name)

        # Save Markdown output
        markdown_path = os.path.join(output_dir, "report.md")
        with open(markdown_path, "w", encoding="utf-8") as f:
            f.write(final_output)

        # Update status
        processing_status[session_id].update({
            "status": "completed",
            "message": "Report generation complete!",
            "progress": 1.0,
            "current_stage": "completed",
            "final_output": final_output,
            "markdown_path": markdown_path
        })

    except Exception as e:
        processing_status[session_id].update({
            "status": "error",
            "message": f"Error generating report: {str(e)}",
            "progress": 0,
            "current_stage": "error"
        })


@app.post("/upload", response_model=ProcessingResponse)
async def upload_pdf(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    """Upload PDF file for processing"""
    # Generate unique session ID
    session_id = str(uuid.uuid4())
    
    # Create file path
    file_path = os.path.join("uploads", f"{session_id}_{file.filename}")
    
    # Save uploaded file
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Start processing in background
    background_tasks.add_task(process_pdf, session_id, file_path, background_tasks)
    
    return {"session_id": session_id, "message": "PDF upload successful. Processing started."}

@app.get("/status/{session_id}", response_model=ProcessingStatus)
async def get_status(session_id: str):
    """Get current processing status"""
    if session_id not in processing_status:
        raise HTTPException(status_code=404, detail="Processing session not found")
    
    status_data = processing_status[session_id]
    
    return {
        "session_id": session_id,
        "status": status_data.get("status", "unknown"),
        "message": status_data.get("message", ""),
        "progress": status_data.get("progress", 0.0),
        "current_stage": status_data.get("current_stage", "unknown"),
        "total_pages": status_data.get("total_pages", 0),
        "processed_pages": status_data.get("processed_pages", 0),
        "final_output": status_data.get("final_output", None)  # Include the final report content
    }

@app.get("/pages/{session_id}/{page_number}", response_model=PageData)
async def get_page_data(session_id: str, page_number: int):
    """Get data for a specific page"""
    if session_id not in processing_status:
        raise HTTPException(status_code=404, detail="Processing session not found")
    
    status_data = processing_status[session_id]
    page_key = f"Page {page_number}"
    
    if page_key not in status_data.get("extracted_pages", {}):
        raise HTTPException(status_code=404, detail=f"Page {page_number} not found")
    
    return {
        "page_number": page_number,
        "raw_text": status_data["extracted_pages"].get(page_key, ""),
        "translated_text": status_data["translated_pages"].get(page_key, "")
    }

@app.get("/image/{session_id}/{page_number}")
async def get_page_image(session_id: str, page_number: int):
    """Get image for a specific page"""
    if session_id not in processing_status:
        raise HTTPException(status_code=404, detail="Processing session not found")
    
    status_data = processing_status[session_id]
    
    if int(page_number)-1 not in status_data.get("pdf_images", {}):
        raise HTTPException(status_code=404, detail=f"Image for page {page_number} not found")
    
    # Return base64 encoded image
    return {"image": status_data["pdf_images"].get(int(page_number)-1, "")}

@app.put("/update-page/{session_id}", response_model=dict)
async def update_page_text(session_id: str, data: PageUpdateRequest):
    """Update edited text for a page"""
    if session_id not in processing_status:
        raise HTTPException(status_code=404, detail="Processing session not found")
    
    page_key = f"Page {data.page_number}"
    processing_status[session_id]["edited_pages"][page_key] = data.edited_text
    
    # Save updated edited pages
    session_dir = os.path.join("temp", session_id)
    with open(os.path.join(session_dir, "edited_pages.json"), "w", encoding="utf-8") as f:
        json.dump(processing_status[session_id]["edited_pages"], f, ensure_ascii=False, indent=2)
    
    return {"status": "success", "message": f"Page {data.page_number} updated successfully"}


@app.post("/generate-report/{session_id}", response_model=dict)
async def start_report_generation(data: ReportRequest, background_tasks: BackgroundTasks):
    """Start report generation process"""
    session_id = data.session_id
    
    if session_id not in processing_status:
        raise HTTPException(status_code=404, detail="Processing session not found")
    
    # Start report generation in background
    background_tasks.add_task(generate_report, session_id, data.client_name)
    
    return {"status": "success", "message": "Report generation started"}



@app.get("/download/{session_id}/{file_type}")
async def download_file(session_id: str, file_type: str):
    """Download generated report file"""
    if session_id not in processing_status:
        raise HTTPException(status_code=404, detail="Processing session not found")
    
    status_data = processing_status[session_id]
    
    if file_type == "markdown":
        final_output = status_data.get("final_output")
        if not final_output:
            raise HTTPException(status_code=404, detail="No markdown content available")
        
        output_dir = os.path.join("outputs", session_id)
        os.makedirs(output_dir, exist_ok=True)
        markdown_path = os.path.join(output_dir, "report.md")
        
        logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../frontend/foxmandal-logo.png"))
        if os.path.exists(logo_path):
            logo_markdown = f"""![](file:///{logo_path}){{width=3in height=1in}}

---

"""
        else:
            logo_markdown = """---

"""
        
        content_with_logo = logo_markdown + final_output
        
        with open(markdown_path, 'w', encoding='utf-8') as f:
            f.write(content_with_logo)
        
        status_data["markdown_path"] = markdown_path
        
        return FileResponse(
            path=markdown_path,
            media_type="text/markdown",
            filename="report.md",
            content_disposition_type="attachment"
        )
        
    elif file_type == "docx":
        final_output = status_data.get("final_output")
        if not final_output:
            raise HTTPException(status_code=404, detail="No content available for conversion")
        
        output_dir = os.path.join("outputs", session_id)
        os.makedirs(output_dir, exist_ok=True)
        docx_path = os.path.join(output_dir, "report.docx")
        
        try:
            from datetime import datetime
            from docx.enum.table import WD_ALIGN_VERTICAL
            
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Aptos'
            font.size = Pt(12)
            
            
            # Add logo at the top of the first page only
            logo_para = doc.add_paragraph()
            logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../frontend/foxmandal-logo.png"))
            if os.path.exists(logo_path):
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_path, width=Inches(2.8), height=Inches(0.9))
                logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                logo_para.space_after = Pt(6)
            
            # Add horizontal line after logo
            line_para = doc.add_paragraph()
            line_run = line_para.add_run("_" * 85)
            line_run.font.color.rgb = RGBColor(64, 64, 64)
            line_run.font.size = Pt(8)
            line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            line_para.space_after = Pt(12)
            
            # Create main content table with better spacing
            first_page_table = doc.add_table(rows=1, cols=2)
            first_page_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Set table width and column widths
            first_page_table.autofit = False
            first_page_table.allow_autofit = False
            first_page_table.columns[0].width = Inches(3.8)  # Left column for address
            first_page_table.columns[1].width = Inches(2.7)  # Right column for title and date
            
            # Remove all table borders for clean look
            tbl = first_page_table._element
            tblPr = tbl.tblPr
            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)
            
            tblBorders = OxmlElement('w:tblBorders')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'none')
                tblBorders.append(border)
            tblPr.append(tblBorders)
            
            # Left cell - Address section
            address_cell = first_page_table.cell(0, 0)
            address_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Clear and format address
            address_para = address_cell.paragraphs[0]
            address_para.clear()
            
            # Add "To," with proper spacing
            to_run = address_para.add_run("To,")
            to_run.font.name = 'Aptos'
            to_run.font.size = Pt(11)
            to_run.font.bold = True
            
            # Add line break
            address_para.add_run("\n\n")
            
            # Add company name
            company_run = address_para.add_run("Vivid Renewables Private Limited,")
            company_run.font.name = 'Aptos'
            company_run.font.size = Pt(11)
            company_run.font.bold = True
            
            # Add address details
            address_details = """\nRegional Office @ Astra Tower, 5th Floor,
Chetan Vihar, Plot No: 15 to 20
Chetan college Road, Shirur Park, Vidyanagar
Hubli- 580021, Karnataka, India."""
            
            address_detail_run = address_para.add_run(address_details)
            address_detail_run.font.name = 'Aptos'
            address_detail_run.font.size = Pt(10)
            
            address_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            address_para.space_after = Pt(6)
            
            # Right cell - Title and Date section
            title_date_cell = first_page_table.cell(0, 1)
            title_date_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Clear and format title/date
            title_date_para = title_date_cell.paragraphs[0]
            title_date_para.clear()
            
            # Add report title
            title_run = title_date_para.add_run("Report on the Title")
            title_run.font.name = 'Aptos'
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title_run.font.italic = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add spacing between title and date
            title_date_para.add_run("\n\n\n")
            
            # Format current date with ordinal
            day = datetime.now().day
            if 4 <= day <= 20 or 24 <= day <= 30:
                suffix = "th"
            else:
                suffix = ["st", "nd", "rd"][day % 10 - 1]
            
            formatted_date = datetime.now().strftime(f"%d{suffix} %B %Y")
            
            date_run = title_date_para.add_run(formatted_date)
            date_run.font.name = 'Aptos'
            date_run.font.size = Pt(12)
            date_run.font.bold = True
            date_run.font.color.rgb = RGBColor(0, 0, 0)
            
            title_date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            title_date_para.space_after = Pt(20)
            
            # Add proper spacing after first page content
            doc.add_paragraph()
            spacer_para = doc.add_paragraph()
            spacer_para.space_after = Pt(12)
            
            def add_heading_with_background(doc, text, level=1):
                # Create a table with one cell to hold the heading
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                
                # Set table width to full page width
                table.autofit = False
                table.allow_autofit = False
                
                # Set column width
                for column in table.columns:
                    column.width = Inches(5.5)
                
                # Get the cell
                cell = table.cell(0, 0)
                
                # Set cell background color to #088484
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Remove any existing shading
                existing_shd = tcPr.find(qn('w:shd'))
                if existing_shd is not None:
                    tcPr.remove(existing_shd)
                
                # Create new shading element with #088484 background
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '088484')  # #088484 background
                tcPr.append(shd)
                
                # Set cell margins for padding
                tcMar = OxmlElement('w:tcMar')
                for side in ['top', 'left', 'bottom', 'right']:
                    margin_elem = OxmlElement(f'w:{side}')
                    margin_elem.set(qn('w:w'), '144')  # 0.15 inch padding
                    margin_elem.set(qn('w:type'), 'dxa')
                    tcMar.append(margin_elem)
                tcPr.append(tcMar)
                
                # Clear existing content and add formatted text
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                
                # Add the text with white color and bold formatting
                run = paragraph.add_run(text)
                run.font.name = 'Aptos'
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                run.font.size = Pt(16) if level == 1 else Pt(14) if level == 2 else Pt(12)  # Adjust size by level
                
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.style = None  # Clear any inherited styles to ensure white text
                
                # Remove all table borders
                tbl = table._element
                tblPr = tbl.tblPr
                existing_borders = tblPr.find(qn('w:tblBorders'))
                if existing_borders is not None:
                    tblPr.remove(existing_borders)
                
                tblBorders = OxmlElement('w:tblBorders')
                for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{side}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tblBorders.append(border)
                tblPr.append(tblBorders)
                
                # Add spacing after the heading
                doc.add_paragraph()
                
                return table
            
            def style_table_header(table):
                if table.rows:
                    header_row = table.rows[0]
                    for cell in header_row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                run.font.bold = True
                                run.font.name = 'Aptos'
                                run.font.size = Pt(11)
            
            lines = final_output.split('\n')
            current_table = None
            table_headers = []
            
            for line in lines:
                line_stripped = line.strip()
                
                if not line_stripped:
                    doc.add_paragraph()
                    continue
                
                if line_stripped.startswith('I. ') or line_stripped.startswith('II. ') or line_stripped.startswith('III. ') or line_stripped.startswith('IV. ') or line_stripped.startswith('V. ') or line_stripped.startswith('VI. '):
                    current_table = None
                    heading_text = line_stripped
                    add_heading_with_background(doc, heading_text, 1)
                elif line_stripped.startswith('### '):
                    current_table = None
                    add_heading_with_background(doc, line_stripped[4:], 3)
                elif line_stripped.startswith('## '):
                    current_table = None
                    add_heading_with_background(doc, line_stripped[3:], 2)
                elif line_stripped.startswith('# '):
                    current_table = None
                    add_heading_with_background(doc, line_stripped[2:], 1)
                elif line_stripped.startswith('| ') and line_stripped.endswith(' |'):
                    cells = [cell.strip() for cell in line_stripped.strip('|').split('|')]
                    
                    if current_table is None:
                        current_table = doc.add_table(rows=1, cols=len(cells))
                        current_table.style = 'Table Grid'
                        current_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                        table_headers = cells
                        
                        header_row = current_table.rows[0]
                        for i, cell_text in enumerate(cells):
                            header_row.cells[i].text = cell_text
                        
                        style_table_header(current_table)
                    else:
                        row = current_table.add_row()
                        for i, cell_text in enumerate(cells):
                            if i < len(row.cells):
                                row.cells[i].text = cell_text
                elif re.match(r'^\|[\s\-\|]+\|$', line_stripped):
                    continue
                elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                    current_table = None
                    doc.add_paragraph(line_stripped[2:], style='List Bullet')
                elif re.match(r'^\d+\.\s', line_stripped):
                    current_table = None
                    doc.add_paragraph(line_stripped[line_stripped.find('.') + 1:].strip(), style='List Number')
                elif '**' in line_stripped:
                    current_table = None
                    paragraph = doc.add_paragraph()
                    parts = line_stripped.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            paragraph.add_run(part)
                        else:
                            run = paragraph.add_run(part)
                            run.bold = True
                else:
                    current_table = None
                    doc.add_paragraph(line_stripped)
            
            doc.save(docx_path)
            status_data["docx_path"] = docx_path
            
            return FileResponse(
                path=docx_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename="report.docx",
                content_disposition_type="attachment"
            )
            
        except Exception as docx_error:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to generate DOCX: {str(docx_error)}"
            )

@app.get("/poor-quality-pages/{session_id}")
async def get_poor_quality_pages(session_id: str):
    """Get list of poor quality pages for a session"""
    try:
        session_dir = os.path.join("temp", session_id)
        poor_quality_file = os.path.join(session_dir, "poor_quality_pages.json")
        
        if not os.path.exists(poor_quality_file):
            return []
            
        with open(poor_quality_file, "r", encoding="utf-8") as f:
            poor_quality_pages = json.load(f)
            
        return poor_quality_pages
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Add caching for document analysis
@lru_cache(maxsize=100)
def analyze_document_text(text: str) -> List[Dict]:
    """
    Analyze document text and return required documents.
    Results are cached for 5 minutes to improve performance.
    """
    # Define document keywords mapping
    document_keywords = {
        "E-Stamp Application": ["estamp", "e-stamp", "stamp"],
        "Form 18": ["form 18", "application for form 18"]
    }
    
    # Convert text to lowercase once
    text_lower = text.lower()
    
    # Find missing documents
    missing_documents = []
    for doc_name, keywords in document_keywords.items():
        if not any(keyword in text_lower for keyword in keywords):
            missing_documents.append({
                "name": doc_name,
                "required": True,
                "uploaded": False
            })
    
    return missing_documents

@app.post("/document-suggestions/")
async def get_document_suggestions(request: DocumentRequest):
    text_lower = request.chunk_text.lower()
    # Remove all whitespace for loose matching
    text_no_space = re.sub(r'\s+', '', text_lower)
    missing_docs = []
    More tolerant E-Stamp detection
    if not (
        "estamp" in text_no_space or
        "e-stamp" in text_lower or
        "stamp duty" in text_lower or
        re.search(r"e[\s\n\r\t]*stamp", text_lower)
    ):
        missing_docs.append({"name": "E-STAMP", "required": True, "uploaded": False})
    if not any(keyword in text_lower for keyword in ["form 18", "application for form 18"]):
        missing_docs.append({"name": "Form 18", "required": True, "uploaded": False})
    return missing_docs

# Mount static files for frontend
app.mount("/", StaticFiles(directory="../frontend/build", html=True), name="frontend")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)